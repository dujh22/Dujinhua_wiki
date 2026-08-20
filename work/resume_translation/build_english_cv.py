from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('work/resume_translation/Jinhua_Du_CV_English.docx')
MEDIA = Path('work/resume_translation/media/word/media')
NAVY = '17365D'
MID = '4C5D73'
LIGHT = 'E8EDF3'
TEXT = RGBColor(25, 31, 40)
GRAY = RGBColor(78, 91, 108)


def set_cell_margins(cell, top=35, start=55, bottom=35, end=55):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, **edges):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge, opts in edges.items():
        tag = f'w:{edge}'
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key in ['val', 'sz', 'space', 'color']:
            if key in opts:
                element.set(qn(f'w:{key}'), str(opts[key]))


def set_run_font(run, name='Arial', size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para(p, before=0, after=0, line=1.0, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep


def add_text(p, text, bold=False, italic=False, size=8.4, color=TEXT):
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return r


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '2F5F91')
    underline = OxmlElement('w:u'); underline.set(qn('w:val'), 'single')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial'); rFonts.set(qn('w:hAnsi'), 'Arial')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '17')
    rPr.extend([color, underline, rFonts, sz])
    new_run.append(rPr)
    text_el = OxmlElement('w:t'); text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def section_heading(doc, title):
    p = doc.add_paragraph()
    set_para(p, before=4, after=3, line=1.0, keep=True)
    add_text(p, title.upper(), bold=True, size=12.2, color=RGBColor(23,54,93))
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr'); pPr.append(pbdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), NAVY)
    pbdr.append(bottom)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None, size=8.2):
    p = doc.add_paragraph()
    set_para(p, before=0, after=1, line=1.0)
    p.paragraph_format.left_indent = Inches(0.16 + 0.16 * level)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    add_text(p, '• ', bold=True, size=size, color=RGBColor(23,54,93))
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, size=size)
        add_text(p, text[len(bold_prefix):], size=size)
    else:
        add_text(p, text, size=size)
    return p


def add_label_line(doc, label, text, size=8.2):
    p = doc.add_paragraph()
    set_para(p, after=1, line=1.0)
    add_text(p, label + ': ', bold=True, size=size, color=RGBColor(23,54,93))
    add_text(p, text, size=size)
    return p


def project_header(doc, role, title, tag=None):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(5.65)
    t.columns[1].width = Inches(1.45)
    left, right = t.rows[0].cells
    shade_cell(left, LIGHT); shade_cell(right, LIGHT)
    set_cell_margins(left, 45, 65, 45, 65); set_cell_margins(right, 45, 65, 45, 65)
    p = left.paragraphs[0]; set_para(p, keep=True)
    add_text(p, role, bold=True, size=8.6, color=RGBColor(23,54,93))
    add_text(p, '  |  ' + title, bold=True, size=8.6)
    p2 = right.paragraphs[0]; set_para(p2, keep=True); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if tag: add_text(p2, tag, italic=True, size=7.7, color=GRAY)
    return t


def add_job(doc, dates, org, role):
    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    t.columns[0].width = Inches(1.25); t.columns[1].width = Inches(4.55); t.columns[2].width = Inches(1.25)
    for c in t.rows[0].cells: set_cell_margins(c, 28, 45, 28, 45)
    p=t.cell(0,0).paragraphs[0]; set_para(p); add_text(p, dates, bold=True, size=8.0, color=RGBColor(23,54,93))
    p=t.cell(0,1).paragraphs[0]; set_para(p); add_text(p, org, size=8.1)
    p=t.cell(0,2).paragraphs[0]; set_para(p); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; add_text(p, role, italic=True, size=7.9, color=GRAY)


def page_break(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
sec.top_margin = Inches(0.42); sec.bottom_margin = Inches(0.42)
sec.left_margin = Inches(0.48); sec.right_margin = Inches(0.48)
sec.header_distance = Inches(0.18); sec.footer_distance = Inches(0.18)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Arial'; normal.font.size = Pt(8.4); normal.font.color.rgb = TEXT
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

# Header
head = doc.add_table(rows=2, cols=2)
head.autofit = False
head.columns[0].width = Inches(5.95); head.columns[1].width = Inches(1.25)
head.cell(0,0).merge(head.cell(1,0))
left = head.cell(0,0); photo = head.cell(0,1); qrs = head.cell(1,1)
for c in (left, photo, qrs): set_cell_margins(c, 20, 30, 20, 30)
left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
p=left.paragraphs[0]; set_para(p, after=1)
add_text(p, 'JINHUA DU', bold=True, size=23, color=RGBColor(23,54,93))
p=left.add_paragraph(); set_para(p, after=1)
add_text(p, 'Ph.D. Candidate in Computer Science and Technology', bold=True, size=10.5)
p=left.add_paragraph(); set_para(p, after=1)
add_text(p, 'Large Language Models  |  Natural Language Processing  |  Machine Learning  |  Big Data', size=8.5, color=GRAY)
p=left.add_paragraph(); set_para(p, after=0)
add_text(p, 'Advisor: Prof. Jie Tang  |  Fourth-year Ph.D. student  |  Expected graduation: 2027-2028', size=8.3)
p=left.add_paragraph(); set_para(p)
add_text(p, 'Homepage: ', bold=True, size=8.3)
add_hyperlink(p, 'https://dujh22.github.io/Dujinhua_wiki/', 'https://dujh22.github.io/Dujinhua_wiki/')
pp=photo.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_para(pp)
pp.add_run().add_picture(str(MEDIA/'image3.jpeg'), width=Inches(0.82), height=Inches(1.15))
qt=qrs.add_table(rows=1, cols=2)
qt.autofit = False
qt.columns[0].width = Inches(0.55); qt.columns[1].width = Inches(0.55)
for i,img in enumerate(['image1.png','image2.png']):
    qc=qt.cell(0,i); set_cell_margins(qc,0,5,0,5)
    qp=qc.paragraphs[0]; qp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_para(qp)
    qp.add_run().add_picture(str(MEDIA/img), width=Inches(0.42), height=Inches(0.42))

section_heading(doc, 'Profile')
info = doc.add_table(rows=2, cols=4)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
vals = [
    ('GPA','3.95 / 4.00'), ('Rank','1 / 329'), ('Gender','Male'), ('Age','25'),
    ('Department','Computer Science'), ('Political affiliation','CPC Member'), ('Role','Graduate Party Branch Secretary'), ('Additional role','Class Advisor / Counselor'),
]
for idx,(lab,val) in enumerate(vals):
    c=info.cell(idx//4,idx%4); shade_cell(c, 'F4F6F8'); set_cell_margins(c,42,55,42,55)
    p=c.paragraphs[0]; set_para(p)
    add_text(p, lab+'\n', bold=True, size=7.4, color=RGBColor(23,54,93)); add_text(p,val,size=7.7)

section_heading(doc, 'Education and Skills')
project_header(doc, 'Sep. 2022 - Jun. 2027', 'Ph.D. in Computer Science and Technology, Tsinghua University', 'Beijing')
add_label_line(doc, 'Selected coursework', 'Advanced Machine Learning (A); Big Data Analysis and Processing (A); Computational Linguistics (A); Frontiers of Information Retrieval (A); Network Computing and Blockchain Technology (A); Digital-Intelligence Security and Standardization (A); Principles and Algorithms of Data Mining; Knowledge Engineering; Principles of Artificial Intelligence; and related courses.', 7.9)
add_label_line(doc, 'Research interests', 'Large language models, natural language processing, machine learning, and big data analytics.', 7.9)
add_label_line(doc, 'Key achievements', '13 papers: 1 CCF-A paper in ML, 1 CCF-A paper in NLP, 2 CCF-C papers, 2 CCF-A papers on LLMs, 2 SCI Q1 papers, and 1 SCI Q3 paper. Trained five large models end-to-end: GLM, LogicGLM, MalayGLM, AiMed, and a public-security LLM.', 7.9)
project_header(doc, 'Sep. 2018 - Jun. 2022', 'B.Eng. in Computer Science and Technology, China University of Geosciences (Beijing)', 'Beijing')
add_label_line(doc, 'Selected coursework', 'C++ Programming (100); Scientific Engineering and Computing (100); High-Performance Computing (99); Artificial Intelligence (97.9); Principles of Operating Systems (97); Data Structures (94); Software Engineering (Excellent); and related courses.', 7.9)
add_label_line(doc, 'Research interests', 'Machine learning, numerical analysis, and robotics (multi-agent systems).', 7.9)
add_label_line(doc, 'Key achievements', '6 papers: 3 on machine learning and numerical computing/modeling (2 EI-indexed and 1 core-journal paper), and 3 on machine learning, intelligent control, and robotics (1 SCI, 1 EI-indexed, and 1 core-journal paper).', 7.9)
add_label_line(doc, 'English proficiency', 'IELTS 7.0; CET-6 525; CET-4 580.', 7.9)
add_label_line(doc, 'Technical skills', 'Proficient in Python and C/C++; familiar with Java, Lua, Go, and other languages; experienced with MATLAB for numerical analysis and computing.', 7.9)

section_heading(doc, 'Honors and Competition Awards')
add_bullet(doc, 'National Scholarship; National Encouragement Scholarship; Outstanding Graduate Student Leader, Tsinghua University; Captain of a Gold Award Student Team; Excellence Award at the Master\'s and Doctoral Forum (First Place).', size=7.9)
add_bullet(doc, 'External Expert for regional public security agencies; Outstanding Undergraduate Graduate of Beijing; Outstanding Undergraduate Thesis of Beijing; "Top Ten Student" of Beidi Pioneer.', size=7.9)
add_bullet(doc, 'National first prizes: MCM/ICM; Gold Award in the main track of the 2024 "Challenge Cup" Capital University Student Entrepreneurship Competition.', size=7.9)
add_bullet(doc, 'National second prizes: National University Computer Skills Challenge; RoboCup China Open.', size=7.9)
add_bullet(doc, 'Provincial/ministerial first prizes: Lanqiao Cup National Software and IT Professionals Competition; Mobile Application Innovation Competition.', size=7.9)
add_bullet(doc, 'More than 50 additional representative competition awards; see the QR codes or personal homepage.', size=7.9)

# Major research (continues naturally)
section_heading(doc, 'Research Experience - Major Projects')
project_header(doc, 'Primary Research Project (Project Lead)', 'Key Technologies for Deep and Complex Reasoning in Foundation Models', 'Doctoral theoretical research')
add_label_line(doc, 'Subproject', 'Complex Reasoning with Foundation Models, under the National Key R&D Program of China (Tsinghua University, Tianjin University, and Zhipu AI).')
add_bullet(doc, 'Paper: "GLM-4.5: Agentic, Reasoning, and Coding (ARC) Foundation Models" - arXiv.')
add_bullet(doc, 'Paper: "A Survey of Post-Training Scaling in Large Language Models" - CCF-A, ACL 2025.')
add_bullet(doc, 'Paper: "Enhancing Test-Time Learning of LLM Agents via RL over Memory" - CCF-A, ICML 2026.')
add_bullet(doc, 'Paper: "Toward Evolvable Evaluation for Logical Reasoning" - under review, CCF-A, NeurIPS.')
add_bullet(doc, 'Paper: "A Self-Evolving Closed-Loop Learning Framework for Logical Reasoning" - CCF-A, submitted to AAAI.')

project_header(doc, 'Primary Research Project (Project Lead)', 'Natural Language Processing for Medical Big Data Analysis and Utilization', 'Doctoral applied research')
add_label_line(doc, 'Subproject', 'Medical Knowledge Graphs Supporting Healthcare Data Elements.')
add_bullet(doc, 'Paper: "Research and Advances in Named Entity Recognition for Chinese Electronic Medical Records" - CCF-A, Acta Electronica Sinica.')
add_bullet(doc, 'Paper: "KrNER: A Novel Named Entity Recognition Method Based on Knowledge Enhancement and Remote Supervision" - CCF-C, TrustCom.')
add_bullet(doc, 'Paper: "KLDP: A Data Profiling Technique Based on Knowledge Graph and LLM" - CCF-C, TrustCom.')
add_label_line(doc, 'Subproject', 'AiMed: Smart Healthcare LLMs Driven by Both Data and Knowledge.')
add_bullet(doc, 'Invited talk: "ChatGPT Has Sparked an AI Craze" - invited by Microsoft.')
add_bullet(doc, 'Paper: "AiMed: Artificial Intelligence Large Language Model for Chinese Medicine" - EI, MedAI; software copyright.')
add_bullet(doc, 'Paper: "Towards Artificial Intelligence for Science: A Case Study of Using ChatGPT" - SCI Q1, Big Data.')
add_bullet(doc, 'Paper: "Large Language Models Driven Reliable Clinical Decision-Making" - Informatics and Health.')
add_bullet(doc, 'Paper: "Toward a Large Language Model-Driven Medical Knowledge Retrieval and QA System: Framework Design and Evaluation" - SCI Q1, Engineering.')
add_label_line(doc, 'Responsibilities', 'Led topic selection, background research, literature review, theoretical innovation, experiments, and paper writing.')

# Page 3 - additional research (continues naturally)
section_heading(doc, 'Research Experience - Additional Projects')
project_header(doc, 'Participating Project (Technical Contributor)', 'Key Technologies for Trusted On-Chain/Off-Chain Data Interaction in Blockchain', 'Ministry of Science and Technology')
add_label_line(doc, 'Subproject', 'Large Language Models and Blockchain.')
add_bullet(doc, 'Website: "Open Data Entry" - public service.')
add_bullet(doc, 'Paper: "Is ChatGPT All That MedNLP Needs? A Systematic Evaluation of Knowledge Discovery Capabilities across Biomedical NLP Tasks" - submitted to Nature.')
add_bullet(doc, 'Paper: "OpenMonet: Open Model Orchestration Network" - in progress.')
add_label_line(doc, 'Responsibilities', 'Conducted interdisciplinary technical research, theoretical innovation, application implementation, and paper writing.')

project_header(doc, 'Participating Project (Technical Contributor)', 'Comprehensive Clinical Assessment, Management, and Early-Warning System for Declining Renal Function in Older Adults', 'Ministry of Science and Technology')
add_label_line(doc, 'Subproject', 'Next-Generation Smart Healthcare through Digital-Intelligence Emergence.')
add_bullet(doc, 'Paper: "Research of Client Selection Algorithm in Cross-Device Federated Learning" - CCF-A, Journal of Software.')
add_bullet(doc, 'Medical Data Annotation System and Chinese Word Segmentation System - related services.')
add_label_line(doc, 'Responsibilities', 'Conducted interdisciplinary medical-engineering research, scenario innovation, application implementation, and paper writing.')

project_header(doc, 'Entrepreneurial Project (Project Lead)', 'Development of a Reinforcement Learning-Based Coaching System for Robot Soccer', 'National-level')
add_bullet(doc, 'Development of a Multi-Level ROS-Based Robot Soccer System - national-level project.')
add_bullet(doc, 'Paper: "Research on the Architecture of Operating Systems Such as ROCOS, a ROS Variant, under Ubuntu" - Chinese Science and Technology Core Journal.')
add_bullet(doc, 'Paper: "System Composition and Optimization of Small Soccer Robot" - English-language EI-indexed.')
add_bullet(doc, 'Paper: "Path Planning Based on an Improved Artificial Potential Field and a Novel Grid" - Peking University Core Journal.')
add_label_line(doc, 'Responsibilities', 'Developed the overall research plan and task allocation; completed core AI algorithm development and system implementation.')

project_header(doc, 'Collaborative Project (Team Lead)', 'Data Modeling and Prediction Studies', None)
add_bullet(doc, 'Paper: "Study on the Effectiveness of a Bottle Ban Based on Principal Component Analysis" - national journal.')
add_bullet(doc, 'Paper: "Global Epidemic Classification Based on the K-Nearest Neighbor Algorithm" - English-language EI-indexed.')
add_bullet(doc, 'Paper: "Mid-Term and Long-Term Prediction of Carbon Emissions in Jiangsu Province Based on PCA-STIRPAT Improved GA-BP" - English-language EI-indexed.')
add_label_line(doc, 'Responsibilities', 'Responsible for AI algorithm development and experiments, literature review, and background research.')

project_header(doc, 'Outsourced Project (Project Lead)', 'Applied AI Systems', 'Software copyrights')
add_bullet(doc, 'AiMed Medical-Knowledge LLM Application Service System - software copyright.')
add_bullet(doc, 'Wind Turbine Aerodynamic Balance Detection System; Intelligent Semi-Automated Hotel System - software copyrights.')
add_label_line(doc, 'Responsibilities', 'Managed project schedules; built system frameworks; developed core programs; delivered solution presentations and defenses.')

# Employment/service (continues naturally)
section_heading(doc, 'Internship Experience')
add_job(doc, 'Mar. 2023 - Present', 'Beijing Zhipu Huazhang Technology Co., Ltd. - AI Institute, China', 'AI R&D Engineer')
add_job(doc, 'Sep. 2021 - Mar. 2023', 'Hunan Wangshu Technology Co., Ltd. - Network Big Data Research Center, China', 'AI R&D Engineer')
add_job(doc, 'Jun. 2021 - Sep. 2021', 'Institute of Information Engineering, Chinese Academy of Sciences', 'Assistant Researcher')
add_job(doc, 'Mar. 2021 - Jun. 2021', 'Alibaba Cloud Native Team, China', 'Algorithm Engineer')
add_job(doc, 'Sep. 2018 - Mar. 2021', 'Information Technology Innovation Center, China University of Geosciences (Beijing)', 'Competition Program Supervisor')
add_label_line(doc, 'Overall period represented in the source', 'Sep. 2018 - Jun. 2022 for the earlier research and industry roles listed above.', 7.8)

section_heading(doc, 'Student Leadership and Service')
add_bullet(doc, 'Teaching assistant for Prof. Jie Tang\'s Advanced Machine Learning. As lead author, responsible for writing and revising the book Large Language Models and establishing its companion website.', size=8.0)
add_bullet(doc, 'Appointed team leader in the Practice Department of the Tsinghua Graduate Youth League Committee; developed and optimized the "Tongxing Practice Platform" WeChat mini program, continuously serving master\'s and doctoral students university-wide.', size=8.0)
add_bullet(doc, 'Published articles on Tsinghua University News as the primary correspondent.', size=8.0)
add_bullet(doc, 'Served as Secretary of the Practice Department, Computer Science Graduate Youth League Committee. As team leader, coordinated publicity and led all joint industry visits across seven departments, including visits to ByteDance, Tencent, Meituan, Ubiquant Investment, NetEase Youdao, and 13 technology companies in Hangzhou; drafted 23 official posts.', size=8.0)
add_bullet(doc, 'Served as a class advisor and received three collective honors: Duxing Model Class/League Branch, Outstanding Class/League Branch (only five university-wide), and Tsinghua University Class-A Youth League Branch.', size=8.0)
add_bullet(doc, 'Received 10 individual honors, including the Second-Class Academic Scholarship, Second-Class Social Work Scholarship, Social Practice Scholarship, Outstanding Communist Youth League Member of Tsinghua University, certificates of service as a Graduate Youth League team leader and departmental secretary, appointment letter as Youth League Branch Secretary, and Outstanding New Student Leader.', size=8.0)
add_bullet(doc, 'Core member of the University Student Union Sports Department; maintained long-term running and fitness training; earned a 10 km race medal, a half-marathon medal, and fifth place in the bodybuilding group.', size=8.0)

section_heading(doc, 'Personal Statement')
p=doc.add_paragraph(); set_para(p, before=1, after=2, line=1.08)
add_text(p, 'Responsible, rigorous, and detail-oriented, with strong logical reasoning and learning ability. I communicate effectively and collaborate well in teams. Through systematic doctoral research training, I hope to create work that genuinely benefits and improves the world, contributing to a kinder, better, and happier future for all.', size=8.6)

# Footer page numbering fields.
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p)
    add_text(p, 'Jinhua Du  •  Curriculum Vitae  •  ', size=7.3, color=GRAY)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rfonts=OxmlElement('w:rFonts'); rfonts.set(qn('w:ascii'),'Arial'); rfonts.set(qn('w:hAnsi'),'Arial')
    color=OxmlElement('w:color'); color.set(qn('w:val'), MID)
    sz=OxmlElement('w:sz'); sz.set(qn('w:val'),'15')
    rPr.extend([rfonts,color,sz]); r.append(rPr)
    t=OxmlElement('w:t'); t.text='1'; r.append(t); fld.append(r); p._p.append(fld)

# Metadata
doc.core_properties.title = 'Jinhua Du - English Curriculum Vitae'
doc.core_properties.subject = 'English translation of curriculum vitae'
doc.core_properties.author = 'Jinhua Du'
doc.core_properties.keywords = 'Jinhua Du, curriculum vitae, computer science, large language models'

doc.save(OUT)
print(OUT)
